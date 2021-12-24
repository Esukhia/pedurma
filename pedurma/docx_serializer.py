import re
from pathlib import Path

from docx import Document


def split_text(content):

    chunks = re.split(r"(\(\d+\) <.*?>)", content)

    return chunks


def create_docx_with_footnotes(text_id, chunks, path):
    document = Document()
    p = document.add_paragraph()

    for chunk in chunks:
        if chunk and "<" in chunk:
            sub_text = p.add_run(chunk)
            sub_text.font.subscript = True
            # sub_text.font.bold = True
            sub_text.font.name = "Jomolhari"
        else:
            normal_text = p.add_run(chunk)
            normal_text.font.name = "Jomolhari"
    output_path = path / f"{text_id}.docx"
    document.save(str(output_path))
    return output_path


def get_pages(text):
    result = []
    pg_text = ""
    pages = re.split(r"(\d+-\d+)", text)
    for i, page in enumerate(pages[:-1]):
        if i % 2 == 0:
            pg_text += page
        else:
            pg_text += page
            result.append(pg_text)
            pg_text = ""
    return result


def parse_page(page, document):
    p = document.add_paragraph()
    chunks = split_text(page)
    for chunk in chunks:
        if chunk and "<" in chunk:
            footnote_number = re.search(r"\((\d+)\) <.*?>", chunk).group(1)
            super_text = p.add_run(f" {footnote_number} ")
            super_text.font.superscript = True
            super_text.font.name = "Jomolhari"
        else:
            normal_text = p.add_run(chunk)
            normal_text.font.name = "Jomolhari"
    notes = re.finditer(r"\((\d+)\) <(.*?)>", page)
    p.add_run("\n---------------\n")
    for note in notes:
        note_text = p.add_run(f"{note.group(1)} {note.group(2)}\n")
        note_text.font.name = "Jomolhari"
    return document


def creat_docx_footnotes_at_end_of_page(text_id, collation_text, path):
    document = Document()
    pages = get_pages(collation_text)
    for page in pages:
        document = parse_page(page, document)
    output_path = path / f"{text_id}.docx"
    document.save(str(output_path))
    return output_path


def get_docx_text(text_id, preview_text, output_path=None, type_="with_footnote"):
    if not output_path:
        (Path.home() / ".collation_docx").mkdir(parents=True, exist_ok=True)
        output_path = Path.home() / ".collation_docx"
    collation_text = ""
    for vol_id, text in preview_text.items():
        collation_text += f"{text}\n\n"
    collation_text = collation_text.replace("\n", "")
    collation_text = re.sub(r"(\d+-\d+)", r"\n\g<1>\n", collation_text)
    if type_ == "with_footnotes":
        chunks = split_text(collation_text)
        docx_path = create_docx_with_footnotes(text_id, chunks, output_path)
    else:
        docx_path = creat_docx_footnotes_at_end_of_page(
            text_id, collation_text, output_path
        )
    return docx_path
